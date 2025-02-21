import os
import docx
from typing import List, Dict, Optional

class DocxLoader:
    """Loader for Microsoft Word documents"""
    
    def __init__(self):
        """Initialize Word document loader"""
        pass
        
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from Word document
        
        Args:
            file_path (str): Path to Word document
            
        Returns:
            str: Extracted text
        """
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
                        
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error extracting text from Word document {file_path}: {str(e)}")
            return ""
    
    def load_documents(self, directory_path: str) -> List[Dict[str, str]]:
        """
        Load all Word documents from a directory
        
        Args:
            directory_path (str): Path to directory containing Word docs
            
        Returns:
            List[Dict[str, str]]: List of documents with metadata
        """
        documents = []
        
        # Check if directory exists
        if not os.path.isdir(directory_path):
            print(f"Directory {directory_path} does not exist")
            return documents
            
        # Process all Word files in the directory
        for filename in os.listdir(directory_path):
            if filename.lower().endswith(('.docx', '.doc')):
                file_path = os.path.join(directory_path, filename)
                try:
                    # Only process .docx files (not legacy .doc)
                    if filename.lower().endswith('.doc'):
                        print(f"Skipping legacy .doc file: {filename}. Convert to .docx format.")
                        continue
                        
                    text = self.extract_text_from_docx(file_path)
                    if text:
                        documents.append({
                            'content': text,
                            'source': file_path,
                            'filename': filename
                        })
                except Exception as e:
                    print(f"Error processing {file_path}: {str(e)}")
                    
        return documents
        
    def load_document(self, file_path: str) -> Optional[Dict[str, str]]:
        """
        Load a single Word document
        
        Args:
            file_path (str): Path to Word document
            
        Returns:
            Optional[Dict[str, str]]: Document with metadata if successful
        """
        if not os.path.isfile(file_path) or not file_path.lower().endswith('.docx'):
            print(f"Invalid Word document: {file_path}")
            return None
            
        try:
            text = self.extract_text_from_docx(file_path)
            if text:
                return {
                    'content': text,
                    'source': file_path,
                    'filename': os.path.basename(file_path)
                }
            return None
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            return None